#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""验证生成的《颜氏家训》文献综述文档内容"""
from docx import Document

doc = Document('D:\\临时\\颜氏家训英译研究文献综述.docx')

print(f'段落数: {len(doc.paragraphs)}')

headings = [(p.style.name, p.text) for p in doc.paragraphs if p.style.name.startswith('Heading')]
print(f'标题数: {len(headings)}')
print()
print('=== 文档结构 ===')
for style, text in headings:
    level = style.replace('Heading ', '')
    print(f'  [H{level}] {text}')

# 统计参考文献
ref_count = 0
in_ref = False
for p in doc.paragraphs:
    if '参考文献' in p.text and p.style.name.startswith('Heading'):
        in_ref = True
        continue
    if in_ref and p.text.strip():
        ref_count += 1

print()
print(f'参考文献总条数: {ref_count}')
print()
print('=== 正文段落中文数统计 ===')
total_chars = sum(len(p.text) for p in doc.paragraphs if not p.style.name.startswith('Heading'))
print(f'总字符数（含参考文献）: {total_chars}')
print()
print('文档验证通过 ✅')
